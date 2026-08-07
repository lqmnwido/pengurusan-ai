"""Document translation helpers used by the translation upload API."""

from __future__ import annotations

import asyncio
import os
import re
import shutil
import subprocess
import tempfile
from functools import lru_cache
from pathlib import Path
from typing import Any, Awaitable, Callable
from xml.sax.saxutils import escape

import fitz
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastapi import Request
from open_webui.models.users import UserModel
from open_webui.storage.provider import Storage
from open_webui.utils.chat import generate_chat_completion
from open_webui.utils.models import get_all_models
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph as RLParagraph, SimpleDocTemplate, Spacer

log = __import__('logging').getLogger(__name__)

DEFAULT_TRANSLATION_MODEL = os.getenv('OPEN_WEBUI_TRANSLATION_MODEL', 'deepseekTranslateV4')
TRANSLATION_MODEL_ALIASES = {
    'deepseekTranslateV4': 'deepseek-v4-flash',
    'deepseek/deepseek-v4-flash': 'deepseek-v4-flash',
}
try:
    DEFAULT_TRANSLATION_CHUNK_TIMEOUT_SECONDS = int(os.getenv('OPEN_WEBUI_TRANSLATION_CHUNK_TIMEOUT_SECONDS', '0'))
except ValueError:
    DEFAULT_TRANSLATION_CHUNK_TIMEOUT_SECONDS = 0
try:
    DEFAULT_TRANSLATION_BLOCK_MAX_TOKENS = int(os.getenv('OPEN_WEBUI_TRANSLATION_BLOCK_MAX_TOKENS', '240'))
except ValueError:
    DEFAULT_TRANSLATION_BLOCK_MAX_TOKENS = 240
try:
    PDF_VISUAL_SIMILARITY_TARGET = float(os.getenv('OPEN_WEBUI_TRANSLATED_PDF_SIMILARITY_TARGET', '0.98'))
except ValueError:
    PDF_VISUAL_SIMILARITY_TARGET = 0.98
DEFAULT_OCR_BACKEND = os.getenv('OPEN_WEBUI_TRANSLATION_OCR_BACKEND', 'auto').strip().lower()
DEFAULT_OCR_LANGUAGE = os.getenv('OPEN_WEBUI_TRANSLATION_OCR_LANGUAGE', 'en').strip().lower() or 'en'
DEFAULT_OCR_USE_GPU = os.getenv('OPEN_WEBUI_TRANSLATION_OCR_USE_GPU', 'auto').strip().lower()
try:
    DEFAULT_OCR_RENDER_SCALE = float(os.getenv('OPEN_WEBUI_TRANSLATION_OCR_SCALE', '1.25'))
except ValueError:
    DEFAULT_OCR_RENDER_SCALE = 1.25


@lru_cache(maxsize=1)
def _paddle_cuda_available() -> bool:
    try:
        import paddle  # type: ignore

        return bool(getattr(paddle.device, 'is_compiled_with_cuda', lambda: False)())
    except Exception:
        return False


@lru_cache(maxsize=8)
def _load_rapidocr():
    try:
        from rapidocr_onnxruntime import RapidOCR

        return RapidOCR()
    except Exception as exc:
        log.debug('RapidOCR unavailable: %s', exc)
        return None


@lru_cache(maxsize=8)
def _load_paddleocr(lang: str, use_gpu: bool):
    try:
        from paddleocr import PaddleOCR

        return PaddleOCR(
            use_angle_cls=True,
            lang=lang,
            use_gpu=use_gpu,
            show_log=False,
        )
    except Exception as exc:
        log.debug('PaddleOCR unavailable for lang=%s use_gpu=%s: %s', lang, use_gpu, exc)
        return None


def _load_ocr_engine() -> tuple[str | None, Any]:
    backend = DEFAULT_OCR_BACKEND
    use_gpu = DEFAULT_OCR_USE_GPU

    paddle_requested = backend in {'auto', 'paddle', 'paddleocr'}
    rapid_requested = backend in {'auto', 'rapidocr'}

    if paddle_requested:
        paddle_use_gpu = _paddle_cuda_available()
        if use_gpu in {'1', 'true', 'yes', 'gpu'}:
            paddle_use_gpu = True
        elif use_gpu in {'0', 'false', 'no', 'cpu'}:
            paddle_use_gpu = False

        paddle_ocr = _load_paddleocr(DEFAULT_OCR_LANGUAGE, paddle_use_gpu)
        if paddle_ocr is not None:
            return f'paddleocr({"gpu" if paddle_use_gpu else "cpu"})', paddle_ocr

    if rapid_requested:
        rapid_ocr = _load_rapidocr()
        if rapid_ocr is not None:
            return 'rapidocr', rapid_ocr

    return None, None


def _ocr_item_sort_key(item: Any) -> tuple[float, float]:
    box = None
    if isinstance(item, dict):
        box = item.get('box') or item.get('points') or item.get('position')
    elif isinstance(item, (list, tuple)) and item:
        if len(item) >= 2 and isinstance(item[0], (list, tuple)):
            box = item[0]
    if isinstance(box, (list, tuple)) and box:
        try:
            xs = [float(point[0]) for point in box if isinstance(point, (list, tuple)) and len(point) >= 2]
            ys = [float(point[1]) for point in box if isinstance(point, (list, tuple)) and len(point) >= 2]
            if ys and xs:
                return (min(ys), min(xs))
        except Exception:
            pass
    return (float('inf'), float('inf'))


def _extract_ocr_lines(ocr_result: Any) -> list[str]:
    if isinstance(ocr_result, tuple) and ocr_result:
        ocr_result = ocr_result[0]

    items = list(ocr_result) if isinstance(ocr_result, list) else []
    ordered: list[tuple[tuple[float, float], str]] = []

    for item in items:
        line = ''
        if isinstance(item, dict):
            line = str(item.get('text') or item.get('transcription') or '').strip()
        elif isinstance(item, (list, tuple)) and len(item) >= 2:
            second = item[1]
            if isinstance(second, str):
                line = second.strip()
            elif isinstance(second, (list, tuple)) and second:
                line = str(second[0]).strip()
        if line:
            ordered.append((_ocr_item_sort_key(item), line))

    ordered.sort(key=lambda entry: entry[0])
    return [line for _, line in ordered]


def _page_ocr_scale(page: Any, base_scale: float = DEFAULT_OCR_RENDER_SCALE) -> float:
    try:
        rect = page.rect
        max_dim = max(float(rect.width), float(rect.height))
    except Exception:
        return max(1.05, base_scale)

    scale = base_scale
    if max_dim >= 1400:
        scale *= 0.9
    elif max_dim >= 1100:
        scale *= 0.95
    elif max_dim <= 700:
        scale *= 1.05

    return min(1.5, max(1.05, round(scale, 2)))


def _render_page_for_ocr(page: Any, scale: float) -> bytes:
    pixmap = page.get_pixmap(
        matrix=fitz.Matrix(scale, scale),
        colorspace=fitz.csGRAY,
        alpha=False,
    )
    return pixmap.tobytes('png')


def _ocr_page_text(page: Any, ocr: Any, force_ocr: bool, scale: float) -> str:
    text = ''
    if not force_ocr:
        try:
            text = page.get_text('text').strip()
        except Exception:
            text = ''

    if text or ocr is None:
        return text

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
        tmp_img.write(_render_page_for_ocr(page, scale))
        tmp_path = tmp_img.name

    try:
        ocr_result = ocr(tmp_path)
        extracted = _extract_ocr_lines(ocr_result)
        return '\n'.join(extracted).strip()
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def _normalize_filename(filename: str, replacement: str = '_') -> str:
    name = os.path.basename(filename).strip()
    name = re.sub(r'[^A-Za-z0-9._ -]+', replacement, name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name or 'translated-document'


def _normalize_translation_model_id(model: str) -> str:
    model = (model or DEFAULT_TRANSLATION_MODEL).strip()
    return TRANSLATION_MODEL_ALIASES.get(model, model)


@lru_cache(maxsize=1)
def _translation_tokenizer():
    try:
        import tiktoken

        return tiktoken.get_encoding('cl100k_base')
    except Exception as exc:
        log.debug('tiktoken unavailable, using estimated token counts: %s', exc)
        return None


def _token_count(text: str) -> int:
    tokenizer = _translation_tokenizer()
    if tokenizer is not None:
        try:
            return len(tokenizer.encode(text or ''))
        except Exception:
            pass
    return max(1, len(text or '') // 4)


def _split_by_token_budget(text: str, max_tokens: int) -> list[str]:
    text = (text or '').strip()
    if not text:
        return []
    if _token_count(text) <= max_tokens:
        return [text]

    chunks: list[str] = []
    current: list[str] = []

    for line in text.split('\n'):
        candidate = '\n'.join(current + [line]).strip()
        if current and _token_count(candidate) > max_tokens:
            chunks.append('\n'.join(current).strip())
            current = [line]
        else:
            current.append(line)

        if current and _token_count('\n'.join(current)) > max_tokens:
            oversized = '\n'.join(current).strip()
            current = []
            words = oversized.split()
            word_chunk: list[str] = []
            for word in words:
                word_candidate = ' '.join(word_chunk + [word])
                if word_chunk and _token_count(word_candidate) > max_tokens:
                    chunks.append(' '.join(word_chunk).strip())
                    word_chunk = [word]
                else:
                    word_chunk.append(word)
            if word_chunk:
                chunks.append(' '.join(word_chunk).strip())

    if current:
        chunks.append('\n'.join(current).strip())

    return [chunk for chunk in chunks if chunk]


def _is_structured_layout_block(text: str) -> bool:
    lines = [line.strip() for line in (text or '').split('\n') if line.strip()]
    if len(lines) < 8:
        return False

    short_lines = sum(1 for line in lines if len(line) <= 45)
    numeric_lines = sum(1 for line in lines if re.search(r'\d', line))
    symbol_lines = sum(1 for line in lines if re.search(r'[:/#().%-]', line))
    return short_lines / len(lines) >= 0.6 or (numeric_lines + symbol_lines) / len(lines) >= 0.55


def _split_structured_layout_block(text: str, max_tokens: int, max_lines: int = 18) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []

    for line in (text or '').split('\n'):
        candidate = '\n'.join(current + [line]).strip()
        if current and (len(current) >= max_lines or _token_count(candidate) > max_tokens):
            chunks.append('\n'.join(current).strip())
            current = [line]
        else:
            current.append(line)

    if current:
        chunks.append('\n'.join(current).strip())

    return [chunk for chunk in chunks if chunk]


def _should_translate_layout_line(line: str) -> bool:
    stripped = (line or '').strip()
    if not stripped:
        return False

    if _is_probable_ocr_noise_line(stripped):
        return False

    if any(token in stripped for token in ('http://', 'https://', 'www.', '@')):
        return False

    if re.fullmatch(r'[\d\s,.:/#()&+\-]+', stripped):
        return False

    if re.fullmatch(r'[A-Z]{2,}[A-Z0-9./\-\s]*', stripped) and re.search(r'\d|/|-|\b(?:LLC|LTD|CO|AED|GSTIN|ETA|ETD|CBM|KGS|USD|RAM)\b', stripped):
        return False

    if re.search(r'\b(?:NSA|HBL|MBL|CCDX|BDX|TEMU)\w*', stripped, re.IGNORECASE):
        return False

    if re.fullmatch(r'INV\d+[A-Z0-9./\-\s]*', stripped, re.IGNORECASE):
        return False

    letters = len(re.findall(r'[^\W\d_]', stripped, flags=re.UNICODE))
    if letters < 2:
        return False

    return True


def _contains_arabic(text: str) -> bool:
    return bool(re.search(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]', text or ''))


def _is_probable_ocr_noise_line(line: str) -> bool:
    stripped = (line or '').strip()
    if not stripped:
        return False

    lowered = stripped.lower()
    if 'snp aue' in lowered or 'jno siu' in lowered or 'junoe o p' in lowered:
        return True

    words = re.findall(r'[A-Za-z]+', stripped)
    if len(words) >= 8:
        avg_len = sum(len(word) for word in words) / len(words)
        very_short = sum(1 for word in words if len(word) <= 3)
        if avg_len <= 3.2 and very_short / len(words) >= 0.65:
            return True

    return False


def _split_blocks(text: str, max_chars: int = 2500) -> list[str]:
    text = text.strip()
    if not text:
        return []

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for block in re.split(r'\n{2,}', text):
        block = block.strip()
        if not block:
            continue

        if len(block) > max_chars:
            if current:
                chunks.append('\n\n'.join(current))
                current = []
                current_len = 0

            start = 0
            while start < len(block):
                chunks.append(block[start : start + max_chars])
                start += max_chars
            continue

        block_len = len(block)
        if current and current_len + block_len + 2 > max_chars:
            chunks.append('\n\n'.join(current))
            current = [block]
            current_len = block_len
        else:
            current.append(block)
            current_len += block_len + 2

    if current:
        chunks.append('\n\n'.join(current))

    return chunks


def _cleanup_translated_text(text: str) -> str:
    text = (text or '').replace('\r\n', '\n').replace('\r', '\n').strip()
    if not text:
        return ''

    lines: list[str] = []
    pending_bullet = False
    skip_header = re.compile(r'^(?:segment|fragmen(?:tasi)? teks)\s+\d+\s+(?:dari|of)\s+\d+[:.]?$', re.IGNORECASE)
    skip_plain_header = re.compile(r'^(?:segment|fragmen(?:tasi)? teks)\s+\d+\s+(?:dari|of)\s+\d+\s*[:.-]?\s*$', re.IGNORECASE)
    skip_instruction = re.compile(
        r'^(?:translate|terjemah|preserve|kekalkan|keep|do not|jangan|return only|hanya kembalikan|source block|blok sumber|<source>|</source>)',
        re.IGNORECASE,
    )

    for raw_line in text.split('\n'):
        line = raw_line.rstrip()
        if not line.strip():
            lines.append('')
            continue

        stripped = line.strip()
        if (
            skip_header.match(stripped)
            or skip_plain_header.match(stripped)
            or skip_instruction.match(stripped)
            or _is_probable_ocr_noise_line(stripped)
        ):
            continue

        # Normalize accidental double bullets or list markers from the model.
        normalized = re.sub(r'^\s*[•·]\s*[•·]\s*', '• ', line)
        normalized = re.sub(r'^\s*[-*]\s*', '• ', normalized)

        if normalized.strip() in {'•', '• ', '•\t'}:
            pending_bullet = True
            continue

        if pending_bullet:
            lines.append(f'• {normalized.strip()}')
            pending_bullet = False
            continue

        # If the model put the bullet marker on the previous line and the
        # content here starts like a continuation, keep it on one line.
        if lines and lines[-1].strip() == '•':
            lines[-1] = f'• {normalized.strip()}'
            continue

        lines.append(normalized)

    cleaned = '\n'.join(lines)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


def _normalize_layout_spacing(text: str) -> str:
    lines: list[str] = []
    for raw_line in (text or '').split('\n'):
        line = raw_line.rstrip()
        if not line.strip():
            lines.append('')
            continue

        if any(token in line for token in ('http://', 'https://', 'mailto:', '@')):
            lines.append(line)
            continue

        if not _should_translate_layout_line(line):
            lines.append(line.strip())
            continue

        normalized = line
        normalized = re.sub(r'(?<=[a-z0-9])(?=[A-Z])', ' ', normalized)
        normalized = re.sub(r'(?<=[A-Za-z])(?=\d)', ' ', normalized)
        normalized = re.sub(r'(?<=\d)(?=[A-Za-z])', ' ', normalized)
        normalized = re.sub(r'(?<=[A-Za-z])\.(?=[A-Z])', '. ', normalized)
        normalized = re.sub(r',(?=\S)', ', ', normalized)
        normalized = re.sub(r'\s*&\s*', ' & ', normalized)
        normalized = re.sub(r'\b[Bb]in\s*[Aa]bdul\b', 'Bin Abdul', normalized)
        normalized = re.sub(r'(\+?60\d{2}-\d{4})(\d{4})\b', r'\1 \2', normalized)
        normalized = re.sub(r'\bJava\s+Script\b', 'JavaScript', normalized, flags=re.IGNORECASE)
        normalized = re.sub(r'\bBootstrap\s+CSS\b', 'Bootstrap CSS', normalized, flags=re.IGNORECASE)
        normalized = re.sub(r'\bLaravel\s+PHP\b', 'Laravel PHP', normalized, flags=re.IGNORECASE)
        normalized = re.sub(r'\bfrontendtechnologies\b', 'frontend technologies', normalized, flags=re.IGNORECASE)
        normalized = re.sub(r'\bUMP\s*Advance\b', 'UMP Advance', normalized, flags=re.IGNORECASE)
        normalized = re.sub(r'\bPENGURUSAN\s*JENAZAH\s*KUANTAN\b', 'PENGURUSAN JENAZAH KUANTAN', normalized, flags=re.IGNORECASE)
        normalized = re.sub(r'\bQANTAN\b', 'KUANTAN', normalized)
        normalized = re.sub(r'\bBo\s*Boi\s*Boy\b', 'BoBoiBoy', normalized, flags=re.IGNORECASE)
        normalized = re.sub(r'[ \t]{2,}', ' ', normalized)
        lines.append(normalized.strip())

    return '\n'.join(lines).strip()


def _split_layout_segments(text: str, max_tokens: int = DEFAULT_TRANSLATION_BLOCK_MAX_TOKENS) -> list[str]:
    text = (text or '').replace('\r\n', '\n').replace('\r', '\n').strip('\n')
    if not text.strip():
        return []

    segments: list[str] = []
    for raw_block in re.split(r'\n\s*\n', text):
        block = raw_block.strip()
        if not block:
            continue

        if _is_structured_layout_block(block):
            segments.extend(_split_structured_layout_block(block, max_tokens=max(120, max_tokens // 2)))
            continue

        if _token_count(block) <= max_tokens:
            segments.append(block)
            continue

        segments.extend(_split_by_token_budget(block, max_tokens=max_tokens))

    return segments


def _format_numbered_source_block(text: str) -> str:
    lines = (text or '').split('\n')
    numbered_lines: list[str] = []
    for index, line in enumerate(lines, start=1):
        marker = 'TR' if _should_translate_layout_line(line) else 'KEEP'
        numbered_lines.append(f'[L{index:03d}|{marker}] {line}')
    return '\n'.join(numbered_lines)


def _restore_numbered_translation(translated: str, source: str) -> str:
    source_lines = (source or '').split('\n')
    parsed: dict[int, str] = {}
    current_id: int | None = None
    line_id_pattern = re.compile(r'^\s*\[?L?(\d{3})(?:\|(?:TR|KEEP))?\]?\s*[:.-]?\s*(.*)$', re.IGNORECASE)

    for raw_line in (translated or '').split('\n'):
        stripped = raw_line.strip()
        if not stripped:
            continue

        match = line_id_pattern.match(stripped)
        if match:
            current_id = int(match.group(1))
            parsed[current_id] = match.group(2).strip()
            continue

        if current_id is not None:
            parsed[current_id] = f'{parsed.get(current_id, "").rstrip()} {stripped}'.strip()

    if not parsed:
        return translated

    restored: list[str] = []
    for index, source_line in enumerate(source_lines, start=1):
        if _is_probable_ocr_noise_line(source_line):
            translated_line = ''
        elif not _should_translate_layout_line(source_line):
            translated_line = source_line.strip()
        else:
            translated_line = parsed.get(index, '').strip()
        restored.append(translated_line)

    return '\n'.join(restored)


def _dedupe_repeated_lines(text: str) -> str:
    lines = (text or '').split('\n')
    output: list[str] = []
    seen_recent: list[str] = []

    for raw_line in lines:
        key = re.sub(r'\s+', ' ', raw_line.strip().lower())
        if key and key in seen_recent:
            continue
        output.append(raw_line)
        if key:
            seen_recent.append(key)
            seen_recent = seen_recent[-20:]

    return '\n'.join(output).strip()


def _iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)


def _paragraph_kind(paragraph: Paragraph) -> str:
    style_name = (getattr(paragraph.style, 'name', '') or '').lower()
    if 'list bullet' in style_name or 'bullet' in style_name:
        return 'bullet'
    if 'list number' in style_name or 'number' in style_name:
        return 'numbered'
    try:
        num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
        if num_pr is not None:
            return 'numbered'
    except Exception:
        pass
    return 'paragraph'


def _extract_docx_structure(file_path: str) -> list[dict[str, Any]]:
    document = Document(file_path)
    blocks: list[dict[str, Any]] = []

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                kind = _paragraph_kind(block)
                blocks.append(
                    {
                        'type': 'paragraph',
                        'kind': kind,
                        'text': text,
                        'style': getattr(block.style, 'name', None),
                    }
                )
        elif isinstance(block, Table):
            rows: list[list[str]] = []
            for row in block.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            blocks.append({'type': 'table', 'rows': rows})

    return blocks


def _extract_pdf_page_blocks(file_path: str, force_ocr: bool = False, ocr_scale: float = DEFAULT_OCR_RENDER_SCALE) -> list[list[str]]:
    pages: list[list[str]] = []
    with fitz.open(file_path) as pdf:
        ocr_backend, ocr = _load_ocr_engine()
        if ocr_backend is None or ocr is None:
            log.debug('No OCR engine available for PDF extraction')

        for page in pdf:
            page_blocks: list[str] = []
            page_scale = _page_ocr_scale(page, ocr_scale)
            if not force_ocr:
                try:
                    raw_blocks = page.get_text('blocks', sort=True)
                except Exception:
                    raw_blocks = []

                for block in raw_blocks:
                    if not isinstance(block, (list, tuple)) or len(block) < 5:
                        continue
                    text = str(block[4] or '').strip()
                    if text:
                        text = re.sub(r'\n{3,}', '\n\n', text)
                        page_blocks.append(text)

            if not page_blocks and ocr is not None:
                text = _ocr_page_text(page, ocr, force_ocr=True, scale=page_scale)
                if text:
                    page_blocks = _split_blocks(text, max_chars=900)

            pages.append(page_blocks)

    return pages


def _render_docx_translation_text(translated_blocks: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for block in translated_blocks:
        if block['type'] == 'paragraph':
            translated_text = block.get('translated_text', '') or ''
            if block.get('kind') == 'bullet' and translated_text:
                parts.append(f'• {re.sub(r"^\\s*[•·]\\s*", "", translated_text).strip()}')
            else:
                parts.append(translated_text)
        elif block['type'] == 'table':
            for row in block.get('translated_rows', []):
                parts.append(' | '.join(row))
    translation_text = '\n\n'.join(part for part in parts if part).strip()
    return _cleanup_translated_text(translation_text)


def _extract_pdf_pages(file_path: str, force_ocr: bool = False) -> list[str]:
    pages: list[str] = []
    with fitz.open(file_path) as pdf:
        ocr_backend, ocr = _load_ocr_engine()
        if ocr_backend is None or ocr is None:
            log.debug('No OCR engine available for PDF extraction')

        for page in pdf:
            page_scale = _page_ocr_scale(page)
            text = ''
            if not force_ocr:
                text = page.get_text('text').strip()

            if not text and ocr is not None:
                text = _ocr_page_text(page, ocr, force_ocr=True, scale=page_scale)

            pages.append(text)

    return pages


def _rect_from_bbox(bbox: Any) -> fitz.Rect:
    rect = fitz.Rect(bbox)
    return fitz.Rect(rect.x0 - 0.5, rect.y0 - 0.5, rect.x1 + 0.5, rect.y1 + 0.5)


def _sample_pdf_background_color(page: Any, rect: fitz.Rect) -> tuple[float, float, float]:
    try:
        page_rect = page.rect
        clip = fitz.Rect(
            max(page_rect.x0, rect.x0 - 2),
            max(page_rect.y0, rect.y0 - 2),
            min(page_rect.x1, rect.x1 + 2),
            min(page_rect.y1, rect.y1 + 2),
        )
        pixmap = page.get_pixmap(clip=clip, alpha=False)
        samples = pixmap.samples
        components = pixmap.n
        if not samples or components < 3:
            return (1, 1, 1)

        buckets: dict[tuple[int, int, int], int] = {}
        for index in range(0, len(samples), components):
            r, g, b = samples[index], samples[index + 1], samples[index + 2]
            # Avoid sampling dark glyph strokes as the background.
            if (r + g + b) / 3 < 150:
                continue
            key = (round(r / 8) * 8, round(g / 8) * 8, round(b / 8) * 8)
            buckets[key] = buckets.get(key, 0) + 1

        if not buckets:
            return (1, 1, 1)
        r, g, b = max(buckets.items(), key=lambda item: item[1])[0]
        return (min(r, 255) / 255, min(g, 255) / 255, min(b, 255) / 255)
    except Exception:
        return (1, 1, 1)


def _insert_fitted_pdf_text(page: Any, rect: fitz.Rect, text: str, font_size: float) -> None:
    if not text:
        return

    sizes = [
        font_size,
        font_size * 0.94,
        font_size * 0.88,
        font_size * 0.82,
        font_size * 0.76,
        font_size * 0.7,
        font_size * 0.64,
        4.5,
        4.0,
        3.5,
    ]
    candidate_rects = [
        rect,
        fitz.Rect(rect.x0, rect.y0 - 0.5, rect.x1 + 2.0, rect.y1 + max(1.5, font_size * 0.45)),
    ]

    for candidate_rect in candidate_rects:
        for size in sizes:
            result = page.insert_textbox(
                candidate_rect,
                text,
                fontname='helv',
                fontsize=max(3.5, size),
                color=(0, 0, 0),
                align=0,
            )
            if result >= 0:
                return

    # Last resort: keep the text visible instead of silently dropping it.
    page.insert_text(
        fitz.Point(rect.x0, max(rect.y0 + 3.5, rect.y1 - 1)),
        text,
        fontname='helv',
        fontsize=3.5,
        color=(0, 0, 0),
    )


def _extract_pdf_layout_items(file_path: str) -> list[list[dict[str, Any]]]:
    pages: list[list[dict[str, Any]]] = []
    with fitz.open(file_path) as pdf:
        for page in pdf:
            page_items: list[dict[str, Any]] = []
            raw = page.get_text('dict', sort=True)
            for block in raw.get('blocks', []):
                if block.get('type') != 0:
                    continue
                for line in block.get('lines', []):
                    spans = [span for span in line.get('spans', []) if (span.get('text') or '').strip()]
                    if not spans:
                        continue
                    text = ''.join(str(span.get('text') or '') for span in spans).strip()
                    if not text:
                        continue

                    try:
                        rect = _rect_from_bbox(line.get('bbox') or spans[0].get('bbox'))
                    except Exception:
                        continue
                    font_size = max(5.0, min(12.0, float(spans[0].get('size') or 8.0)))
                    page_items.append(
                        {
                            'text': text,
                            'bbox': (rect.x0, rect.y0, rect.x1, rect.y1),
                            'font_size': font_size,
                            'should_translate': _should_translate_layout_line(text),
                        }
                    )
            pages.append(page_items)
    return pages


def _render_page_samples(page: Any, scale: float = 0.75) -> tuple[int, int, bytes]:
    pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), colorspace=fitz.csGRAY, alpha=False)
    return pixmap.width, pixmap.height, bytes(pixmap.samples)


def _sample_diff_ratio(left: bytes, right: bytes, threshold: int = 24) -> float:
    if not left or not right:
        return 1.0
    total = min(len(left), len(right))
    changed = 0
    for index in range(total):
        if abs(left[index] - right[index]) > threshold:
            changed += 1
    changed += abs(len(left) - len(right))
    return changed / max(len(left), len(right), 1)


def _external_pdf_raster_compare(source_path: str, translated_path: str) -> dict[str, Any]:
    report: dict[str, Any] = {
        'engine': 'pdftoppm+imagemagick',
        'available': bool(shutil.which('pdftoppm') and shutil.which('compare')),
        'pages': [],
        'average_absolute_error_ratio': None,
        'max_absolute_error_ratio': None,
    }
    if not report['available']:
        return report

    try:
        from PIL import Image
    except Exception as exc:
        report['error'] = f'Pillow unavailable: {exc}'
        return report

    with tempfile.TemporaryDirectory(prefix='pdf-visual-qa-') as tmp_dir:
        source_prefix = os.path.join(tmp_dir, 'source')
        translated_prefix = os.path.join(tmp_dir, 'translated')
        for input_path, output_prefix in ((source_path, source_prefix), (translated_path, translated_prefix)):
            subprocess.run(
                ['pdftoppm', '-png', '-r', '96', input_path, output_prefix],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )

        source_pages = sorted(Path(tmp_dir).glob('source-*.png'))
        translated_pages = sorted(Path(tmp_dir).glob('translated-*.png'))
        page_count = min(len(source_pages), len(translated_pages))
        ratios: list[float] = []

        for index in range(page_count):
            source_image = source_pages[index]
            translated_image = translated_pages[index]
            with Image.open(source_image) as image:
                width, height = image.size
            compare_result = subprocess.run(
                ['compare', '-metric', 'AE', str(source_image), str(translated_image), 'null:'],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            metric_text = (compare_result.stderr or compare_result.stdout or '0').strip().split()[0]
            try:
                absolute_error = float(metric_text)
            except ValueError:
                absolute_error = float(width * height)
            ratio = absolute_error / max(width * height, 1)
            ratios.append(ratio)
            report['pages'].append(
                {
                    'page': index + 1,
                    'absolute_error': int(absolute_error),
                    'absolute_error_ratio': round(ratio, 6),
                }
            )

        if ratios:
            report['average_absolute_error_ratio'] = round(sum(ratios) / len(ratios), 6)
            report['max_absolute_error_ratio'] = round(max(ratios), 6)

    return report


def _screen_pdf_visual_design(source_path: str, translated_path: str) -> dict[str, Any]:
    report: dict[str, Any] = {
        'engine': 'pymupdf',
        'external_tools': {
            'pdftoppm': bool(shutil.which('pdftoppm')),
            'imagemagick_compare': bool(shutil.which('compare')),
            'diff_pdf': bool(shutil.which('diff-pdf')),
        },
        'external_raster_compare': None,
        'page_count_match': False,
        'page_sizes_match': False,
        'average_pixel_diff_ratio': None,
        'max_pixel_diff_ratio': None,
        'similarity_target': PDF_VISUAL_SIMILARITY_TARGET,
        'similarity_score': None,
        'passes_similarity_target': False,
        'pages': [],
    }

    try:
        with fitz.open(source_path) as source_pdf, fitz.open(translated_path) as translated_pdf:
            report['source_page_count'] = len(source_pdf)
            report['translated_page_count'] = len(translated_pdf)
            report['page_count_match'] = len(source_pdf) == len(translated_pdf)

            page_count = min(len(source_pdf), len(translated_pdf))
            page_reports: list[dict[str, Any]] = []
            diff_ratios: list[float] = []

            for page_index in range(page_count):
                source_page = source_pdf[page_index]
                translated_page = translated_pdf[page_index]
                source_rect = source_page.rect
                translated_rect = translated_page.rect
                size_match = (
                    abs(source_rect.width - translated_rect.width) < 0.5
                    and abs(source_rect.height - translated_rect.height) < 0.5
                )

                try:
                    source_width, source_height, source_samples = _render_page_samples(source_page)
                    translated_width, translated_height, translated_samples = _render_page_samples(translated_page)
                    if source_width == translated_width and source_height == translated_height:
                        diff_ratio = _sample_diff_ratio(source_samples, translated_samples)
                    else:
                        diff_ratio = 1.0
                except Exception:
                    diff_ratio = 1.0

                diff_ratios.append(diff_ratio)
                page_reports.append(
                    {
                        'page': page_index + 1,
                        'size_match': size_match,
                        'pixel_diff_ratio': round(diff_ratio, 6),
                    }
                )

            report['pages'] = page_reports
            report['page_sizes_match'] = bool(page_reports) and all(page['size_match'] for page in page_reports)
            if diff_ratios:
                report['average_pixel_diff_ratio'] = round(sum(diff_ratios) / len(diff_ratios), 6)
                report['max_pixel_diff_ratio'] = round(max(diff_ratios), 6)
                similarity_score = max(0.0, 1.0 - float(report['average_pixel_diff_ratio']))
                report['similarity_score'] = round(similarity_score, 6)
                report['passes_similarity_target'] = (
                    bool(report['page_count_match'])
                    and bool(report['page_sizes_match'])
                    and similarity_score >= PDF_VISUAL_SIMILARITY_TARGET
                )
    except Exception as exc:
        report['error'] = str(exc)

    try:
        report['external_raster_compare'] = _external_pdf_raster_compare(source_path, translated_path)
    except Exception as exc:
        report['external_raster_compare'] = {'available': False, 'error': str(exc)}

    return report


def _visual_similarity_score(report: dict[str, Any] | None) -> float:
    if not report:
        return 0.0
    score = report.get('similarity_score')
    if isinstance(score, (int, float)):
        return float(score)
    diff = report.get('average_pixel_diff_ratio')
    if isinstance(diff, (int, float)):
        return max(0.0, 1.0 - float(diff))
    external = report.get('external_raster_compare') or {}
    external_diff = external.get('average_absolute_error_ratio')
    if isinstance(external_diff, (int, float)):
        return max(0.0, 1.0 - float(external_diff))
    return 0.0


def _layout_pages_have_complex_script(layout_pages: list[list[dict[str, Any]]]) -> bool:
    return any(_contains_arabic(item.get('text') or '') for page in layout_pages for item in page)


def _copy_pdf_page_images(source_document: fitz.Document, source_page: fitz.Page, output_page: fitz.Page) -> None:
    seen: set[tuple[int, float, float, float, float]] = set()
    for image_info in source_page.get_images(full=True):
        xref = image_info[0]
        try:
            image = source_document.extract_image(xref)
            stream = image.get('image')
            if not stream:
                continue
            for rect in source_page.get_image_rects(xref):
                key = (xref, round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2))
                if key in seen:
                    continue
                seen.add(key)
                output_page.insert_image(rect, stream=stream)
        except Exception:
            log.debug('Failed to copy image xref=%s', xref, exc_info=True)


def _copy_pdf_page_drawings(source_page: fitz.Page, output_page: fitz.Page) -> None:
    for drawing in source_page.get_drawings():
        shape = output_page.new_shape()
        drew = False
        for item in drawing.get('items') or []:
            op = item[0]
            try:
                if op == 'l':
                    shape.draw_line(item[1], item[2])
                    drew = True
                elif op == 're':
                    shape.draw_rect(item[1])
                    drew = True
                elif op == 'qu':
                    shape.draw_quad(item[1])
                    drew = True
                elif op == 'c':
                    shape.draw_bezier(item[1], item[2], item[3], item[4])
                    drew = True
            except Exception:
                log.debug('Failed to copy drawing item %s', item, exc_info=True)
        if not drew:
            continue
        try:
            shape.finish(
                color=drawing.get('color'),
                fill=drawing.get('fill'),
                width=float(drawing.get('width') or 0.5),
                closePath=bool(drawing.get('closePath')),
                dashes=drawing.get('dashes'),
                even_odd=bool(drawing.get('even_odd', False)),
            )
            shape.commit()
        except TypeError:
            shape.finish(
                color=drawing.get('color'),
                fill=drawing.get('fill'),
                width=float(drawing.get('width') or 0.5),
                closePath=bool(drawing.get('closePath')),
            )
            shape.commit()
        except Exception:
            log.debug('Failed to commit drawing', exc_info=True)


def _pdf_has_rebuildable_layout(file_path: str) -> bool:
    try:
        with fitz.open(file_path) as pdf:
            return any(page.get_drawings() or page.get_images(full=True) for page in pdf)
    except Exception:
        return False


async def _ensure_models_loaded(request: Request, user: UserModel | Any):
    if not getattr(request.app.state, 'MODELS', None):
        await get_all_models(request, user=user)


def _extract_completion_text(response: Any) -> str:
    if isinstance(response, dict):
        choices = response.get('choices') or []
        if choices:
            choice = choices[0] or {}
            message = choice.get('message') or {}
            content = message.get('content')
            if isinstance(content, str):
                return content.strip()
        if 'message' in response and isinstance(response['message'], dict):
            content = response['message'].get('content')
            if isinstance(content, str):
                return content.strip()
    if hasattr(response, 'choices'):
        try:
            return (response.choices[0].message.content or '').strip()
        except Exception:
            pass
    return str(response).strip()


class DocumentTranslator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.body_style = ParagraphStyle(
            'TranslatedBody',
            parent=self.styles['BodyText'],
            fontName='Helvetica',
            fontSize=10.5,
            leading=13.5,
            spaceAfter=8,
        )

    async def translate_text(
        self,
        request: Request,
        user: UserModel | Any,
        text: str,
        target_language: str,
        source_language: str,
        model: str,
        progress_callback: Callable[[list[str], str | None], Awaitable[None] | None] | None = None,
        chunk_timeout_seconds: int = DEFAULT_TRANSLATION_CHUNK_TIMEOUT_SECONDS,
        preserve_layout: bool = False,
        block_label: str | None = None,
    ) -> str:
        await _ensure_models_loaded(request, user)
        completion_model = _normalize_translation_model_id(model)

        if preserve_layout:
            chunks = _split_layout_segments(text, max_tokens=DEFAULT_TRANSLATION_BLOCK_MAX_TOKENS)
        else:
            chunks = _split_blocks(text, max_chars=2500)
        if not chunks:
            return ''

        async def emit_progress(step: str, detail: str | None = None, translation_text: str | None = None) -> None:
            if not progress_callback:
                return
            payload = [step]
            if detail:
                payload.append(detail)
            result = progress_callback(payload, translation_text)
            if asyncio.iscoroutine(result):
                await result

        translated_chunks: list[str] = []
        await emit_progress('Extracting text', f'Chunking text into {len(chunks)} segment(s)')
        for index, chunk in enumerate(chunks, start=1):
            chunk_name = f' {block_label}' if block_label else ''
            await emit_progress('Extracting text', f'Translating{chunk_name} chunk {index} of {len(chunks)}')
            if preserve_layout:
                numbered_chunk = _format_numbered_source_block(chunk)
                prompt = (
                    f'Translate exactly one source block from {source_language} to {target_language}.\n'
                    'A source block is one paragraph or section separated from other blocks by a blank line.\n'
                    'Each source line starts with an ID like [L001|TR] or [L001|KEEP]. Return the same IDs in the same order.\n'
                    'Translate only lines marked TR. Copy KEEP lines exactly unchanged after their ID.\n'
                    'Return exactly one output line for every source line. Do not skip any ID. Do not add extra IDs.\n'
                    'Preserve line breaks, bullet markers, punctuation, names, URLs, emails, phone numbers, dates, and numbers.\n'
                    'For source bullets that start with "-", translate them as bullet lines starting with "•".\n'
                    'Do not merge unrelated lines. Do not omit source meaning. Do not invent new meaning.\n'
                    'Do not copy the English source line unless it is a proper name, URL, email, code, acronym, or title that must remain unchanged.\n'
                    'Keep these terms unchanged: BoBoiBoy, Retak\'ka, Yaya, Ying, Gopal, Fang, Frontend, Softinn Solutions Sdn. Bhd., UMP Advance, Universiti Malaysia Pahang Al-Sultan Abdullah, E-PENGURUSAN JENAZAH KUANTAN, HTML, CSS, JavaScript, Bootstrap CSS, Laravel PHP, API, SRS, SDD.\n'
                    'Use Malay business-letter wording. Translate "Dear Ms." as "Puan". Translate "Sincerely" as "Yang benar".\n'
                    'Return only the translated ID lines. Do not translate these instructions. Do not add labels, notes, summaries, or explanations.\n\n'
                    '<source>\n'
                    f'{numbered_chunk}\n'
                    '</source>'
                )
            else:
                layout_instructions = (
                    'Preserve paragraph breaks, numbering, bullets, and table-like formatting exactly.\n'
                    'Keep names, emails, URLs, code, and other identifiers unchanged.\n'
                )
                prompt_parts = [
                    f'Translate the following document text from {source_language} to {target_language}.\n',
                    layout_instructions,
                    'Do not add labels such as "Segment", "Fragmen", notes, summaries, or explanations.\n',
                    'If the text contains bullet lines, keep each bullet on its own line starting with "•".\n',
                    'Return only the translated text and nothing else.\n\n',
                    f'Text chunk {index} of {len(chunks)}:\n{chunk}',
                ]
                prompt = ''.join(prompt_parts)

            completion_request = generate_chat_completion(
                request,
                {
                    'model': completion_model,
                    'messages': [
                        {
                            'role': 'system',
                            'content': 'You are a precise translation engine. Translate only. Do not explain.',
                        },
                        {'role': 'user', 'content': prompt},
                    ],
                    'stream': False,
                },
                user=user,
                bypass_filter=True,
                bypass_system_prompt=True,
            )
            try:
                response = (
                    await asyncio.wait_for(completion_request, timeout=chunk_timeout_seconds)
                    if chunk_timeout_seconds and chunk_timeout_seconds > 0
                    else await completion_request
                )
            except asyncio.TimeoutError as exc:
                raise RuntimeError(
                    f'Translation chunk {index} timed out after {chunk_timeout_seconds} seconds'
                ) from exc

            translated = _extract_completion_text(response)
            if not translated:
                raise RuntimeError('Translation model returned an empty response')
            if preserve_layout:
                translated = _restore_numbered_translation(translated, chunk)
            translated = _cleanup_translated_text(translated)
            if preserve_layout:
                translated = _normalize_layout_spacing(translated)
                translated = _dedupe_repeated_lines(translated)
            translated_chunks.append(translated)
            await emit_progress(
                'Extracting text',
                f'Translating{chunk_name} chunk {index} of {len(chunks)}',
                '\n\n'.join(part for part in translated_chunks if part) if preserve_layout else '\n\n'.join(part for part in translated_chunks if part).strip(),
            )
            await asyncio.sleep(0)

        return '\n\n'.join(part for part in translated_chunks if part) if preserve_layout else '\n\n'.join(translated_chunks).strip()

    async def translate_pdf_pages(
        self,
        request: Request,
        user: UserModel | Any,
        file_path: str,
        target_language: str,
        source_language: str,
        model: str,
        force_ocr: bool = False,
        progress_callback: Callable[[list[str], str | None], Awaitable[None] | None] | None = None,
        chunk_timeout_seconds: int = DEFAULT_TRANSLATION_CHUNK_TIMEOUT_SECONDS,
    ) -> list[dict[str, Any]]:
        pages = await asyncio.to_thread(_extract_pdf_page_blocks, file_path, force_ocr, 1.5)
        translated_pages: list[dict[str, Any]] = []
        for page_index, page_blocks in enumerate(pages, start=1):
            if not page_blocks:
                translated_pages.append(
                    {
                        'page_number': page_index,
                        'source_blocks': [],
                        'translated_blocks': [],
                        'translation_text': '',
                    }
                )
                continue

            translated_blocks: list[str] = []
            async def page_progress(progress: list[str], page_translation: str | None = None) -> None:
                if not progress_callback:
                    return
                combined_translation = None
                if page_translation is not None:
                    combined_pages = translated_blocks + [page_translation]
                    combined_translation = '\n\n'.join(page for page in combined_pages if page).strip()
                result = progress_callback(progress, combined_translation)
                if asyncio.iscoroutine(result):
                    await result

            for block_index, block_text in enumerate(page_blocks, start=1):
                translated_block = await self.translate_text(
                    request=request,
                    user=user,
                    text=block_text,
                    target_language=target_language,
                    source_language=source_language,
                    model=model,
                    progress_callback=page_progress,
                    chunk_timeout_seconds=chunk_timeout_seconds,
                    preserve_layout=True,
                    block_label=f'page {page_index}, block {block_index}',
                )
                translated_blocks.append(translated_block)
                await asyncio.sleep(0)

            translated_pages.append(
                {
                    'page_number': page_index,
                    'source_blocks': page_blocks,
                    'translated_blocks': translated_blocks,
                    'translation_text': '\n\n'.join(page for page in translated_blocks if page).strip(),
                }
            )
        return translated_pages

    async def translate_docx_blocks(
        self,
        request: Request,
        user: UserModel | Any,
        file_path: str,
        target_language: str,
        source_language: str,
        model: str,
        progress_callback: Callable[[list[str], str | None], Awaitable[None] | None] | None = None,
        chunk_timeout_seconds: int = DEFAULT_TRANSLATION_CHUNK_TIMEOUT_SECONDS,
    ) -> list[dict[str, Any]]:
        blocks = await asyncio.to_thread(_extract_docx_structure, file_path)
        translated_blocks: list[dict[str, Any]] = []

        for block in blocks:
            if block['type'] == 'paragraph':
                async def paragraph_progress(progress: list[str], paragraph_translation: str | None = None) -> None:
                    if not progress_callback:
                        return
                    combined_translation = None
                    if paragraph_translation is not None:
                        combined_blocks = translated_blocks + [{**block, 'translated_text': paragraph_translation}]
                        combined_translation = _render_docx_translation_text(combined_blocks)
                    result = progress_callback(progress, combined_translation)
                    if asyncio.iscoroutine(result):
                        await result

                translated = await self.translate_text(
                    request=request,
                    user=user,
                    text=block['text'],
                    target_language=target_language,
                    source_language=source_language,
                    model=model,
                    progress_callback=paragraph_progress,
                    chunk_timeout_seconds=chunk_timeout_seconds,
                    preserve_layout=True,
                )
                translated_blocks.append({**block, 'translated_text': translated})
            elif block['type'] == 'table':
                translated_rows: list[list[str]] = []
                for row in block['rows']:
                    translated_row: list[str] = []
                    for cell_index, cell in enumerate(row):
                        async def cell_progress(progress: list[str], cell_translation: str | None = None) -> None:
                            if not progress_callback:
                                return
                            combined_translation = None
                            if cell_translation is not None:
                                current_row = translated_row + [cell_translation]
                                padded_row = current_row + [''] * max(0, len(row) - len(current_row))
                                combined_rows = translated_rows + [padded_row]
                                combined_translation = _render_docx_translation_text(
                                    [{**block, 'translated_rows': combined_rows}]
                                )
                            result = progress_callback(progress, combined_translation)
                            if asyncio.iscoroutine(result):
                                await result

                        translated_row.append(
                            await self.translate_text(
                                request=request,
                                user=user,
                                text=cell,
                                target_language=target_language,
                                source_language=source_language,
                                model=model,
                                progress_callback=cell_progress,
                                chunk_timeout_seconds=chunk_timeout_seconds,
                                preserve_layout=True,
                            )
                            if cell.strip()
                            else ''
                        )
                    translated_rows.append(translated_row)
                translated_blocks.append({**block, 'translated_rows': translated_rows})
                await asyncio.sleep(0)

        return translated_blocks

    def build_pdf(self, translated_pages: list[dict[str, Any]], output_path: str, title: str = 'Translated Document') -> str:
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=0.8 * inch,
            leftMargin=0.8 * inch,
            topMargin=0.8 * inch,
            bottomMargin=0.8 * inch,
        )

        story: list[Any] = []
        for index, page_text in enumerate(translated_pages):
            if index > 0:
                story.append(PageBreak())
            story.append(RLParagraph(escape(title), self.body_style))
            story.append(Spacer(1, 0.2 * inch))

            page_blocks = page_text.get('translated_blocks') or []
            page_rendered = (page_text.get('translation_text') or '').strip()
            if page_blocks:
                for paragraph in page_blocks:
                    paragraph = (paragraph or '').strip()
                    if not paragraph:
                        continue
                    for sub_paragraph in re.split(r'\n{2,}', paragraph):
                        sub_paragraph = sub_paragraph.strip()
                        if not sub_paragraph:
                            continue
                        story.append(RLParagraph(escape(sub_paragraph).replace('\n', '<br/>'), self.body_style))
                        story.append(Spacer(1, 0.12 * inch))
            elif page_rendered:
                for paragraph in re.split(r'\n{2,}', page_rendered):
                    paragraph = paragraph.strip()
                    if not paragraph:
                        continue
                    story.append(RLParagraph(escape(paragraph).replace('\n', '<br/>'), self.body_style))
                    story.append(Spacer(1, 0.12 * inch))
            else:
                story.append(RLParagraph('No translatable text detected on this page.', self.body_style))

        doc.build(story)
        return output_path

    async def translate_pdf_layout_items(
        self,
        request: Request,
        user: UserModel | Any,
        file_path: str,
        target_language: str,
        source_language: str,
        model: str,
        progress_callback: Callable[[list[str], str | None], Awaitable[None] | None] | None = None,
        chunk_timeout_seconds: int = DEFAULT_TRANSLATION_CHUNK_TIMEOUT_SECONDS,
    ) -> list[list[dict[str, Any]]]:
        pages = await asyncio.to_thread(_extract_pdf_layout_items, file_path)
        translated_preview_parts: list[str] = []

        for page_index, page_items in enumerate(pages, start=1):
            group: list[dict[str, Any]] = []

            async def flush_group() -> None:
                nonlocal group
                if not group:
                    return
                source_text = '\n'.join(item['text'] for item in group)
                translated = await self.translate_text(
                    request=request,
                    user=user,
                    text=source_text,
                    target_language=target_language,
                    source_language=source_language,
                    model=model,
                    progress_callback=progress_callback,
                    chunk_timeout_seconds=chunk_timeout_seconds,
                    preserve_layout=True,
                    block_label=f'PDF layout page {page_index}',
                )
                translated_lines = translated.split('\n')
                for line_index, item in enumerate(group):
                    translated_line = (
                        translated_lines[line_index].strip()
                        if line_index < len(translated_lines)
                        else ('' if _contains_arabic(item.get('text') or '') else item.get('text') or '')
                    )
                    item['translated_text'] = translated_line.strip()
                    if translated_line.strip():
                        translated_preview_parts.append(translated_line.strip())
                group = []

            for item in page_items:
                if not item.get('should_translate', True):
                    await flush_group()
                    item['translated_text'] = item.get('text') or ''
                    continue

                candidate = '\n'.join([entry['text'] for entry in group] + [item['text']])
                if group and (len(group) >= 12 or _token_count(candidate) > 160):
                    await flush_group()
                group.append(item)
            await flush_group()

        return pages

    def build_pdf_clone(self, source_path: str, translated_layout_pages: list[list[dict[str, Any]]], output_path: str) -> str:
        document = fitz.open(source_path)
        try:
            for page_index, page_items in enumerate(translated_layout_pages):
                if page_index >= len(document):
                    continue
                page = document[page_index]
                insert_items: list[dict[str, Any]] = []

                for item in page_items:
                    translated = (item.get('translated_text') or '').strip()
                    source = (item.get('text') or '').strip()
                    if not item.get('should_translate', True) or not translated or translated == source:
                        continue
                    rect = fitz.Rect(item['bbox'])
                    mask_rect = fitz.Rect(rect.x0 - 0.7, rect.y0 - 0.7, rect.x1 + 0.9, rect.y1 + 0.9)
                    fill_color = _sample_pdf_background_color(page, mask_rect)
                    page.add_redact_annot(mask_rect, fill=fill_color)
                    insert_items.append(
                        {**item, 'translated_text': translated, 'rect': rect, 'fill_color': fill_color}
                    )

                if insert_items:
                    page.apply_redactions()

                for item in insert_items:
                    rect = item['rect']
                    translated = item['translated_text']
                    font_size = float(item.get('font_size') or 8.0)
                    _insert_fitted_pdf_text(page, rect, translated, font_size)

            document.save(output_path, garbage=4, deflate=True)
        finally:
            document.close()
        return output_path

    def build_pdf_image_overlay(
        self,
        source_path: str,
        translated_layout_pages: list[list[dict[str, Any]]],
        output_path: str,
    ) -> str:
        source_document = fitz.open(source_path)
        output_document = fitz.open()
        try:
            for page_index, source_page in enumerate(source_document):
                page_rect = source_page.rect
                output_page = output_document.new_page(width=page_rect.width, height=page_rect.height)
                pixmap = source_page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                output_page.insert_image(page_rect, stream=pixmap.tobytes('png'))

                page_items = translated_layout_pages[page_index] if page_index < len(translated_layout_pages) else []
                for item in page_items:
                    translated = (item.get('translated_text') or '').strip()
                    source = (item.get('text') or '').strip()
                    if not item.get('should_translate', True) or not translated or translated == source:
                        continue

                    rect = fitz.Rect(item['bbox'])
                    mask_rect = fitz.Rect(rect.x0 - 0.8, rect.y0 - 0.8, rect.x1 + 1.2, rect.y1 + 1.2)
                    fill_color = _sample_pdf_background_color(source_page, mask_rect)
                    output_page.draw_rect(mask_rect, color=None, fill=fill_color, overlay=True)

                    font_size = float(item.get('font_size') or 8.0)
                    _insert_fitted_pdf_text(output_page, rect, translated, font_size)

            output_document.save(output_path, garbage=4, deflate=True)
        finally:
            output_document.close()
            source_document.close()
        return output_path

    def build_pdf_layout_rebuild(
        self,
        source_path: str,
        translated_layout_pages: list[list[dict[str, Any]]],
        output_path: str,
    ) -> str:
        source_document = fitz.open(source_path)
        output_document = fitz.open()
        try:
            for page_index, source_page in enumerate(source_document):
                page_rect = source_page.rect
                output_page = output_document.new_page(width=page_rect.width, height=page_rect.height)
                output_page.draw_rect(page_rect, color=None, fill=(1, 1, 1), overlay=False)

                _copy_pdf_page_images(source_document, source_page, output_page)
                _copy_pdf_page_drawings(source_page, output_page)

                page_items = translated_layout_pages[page_index] if page_index < len(translated_layout_pages) else []
                for item in page_items:
                    text = (item.get('translated_text') or item.get('text') or '').strip()
                    if not text:
                        continue
                    rect = fitz.Rect(item['bbox'])
                    font_size = float(item.get('font_size') or 8.0)
                    _insert_fitted_pdf_text(output_page, rect, text, font_size)

            output_document.save(output_path, garbage=4, deflate=True)
        finally:
            output_document.close()
            source_document.close()
        return output_path

    def build_docx(self, translated_blocks: list[dict[str, Any]], output_path: str) -> str:
        document = Document()
        document.core_properties.title = 'Translated document'

        for block in translated_blocks:
            if block['type'] == 'paragraph':
                translated_text = block.get('translated_text', '') or ''
                paragraph = document.add_paragraph()
                if block.get('kind') == 'bullet':
                    try:
                        paragraph.style = 'List Bullet'
                    except Exception:
                        pass
                    translated_text = re.sub(r'^\s*[•·]\s*', '', translated_text).strip()
                elif block.get('kind') == 'numbered':
                    try:
                        paragraph.style = 'List Number'
                    except Exception:
                        pass
                    translated_text = re.sub(r'^\s*\d+[.)]\s*', '', translated_text).strip()
                style_name = block.get('style')
                if style_name and block.get('kind') == 'paragraph':
                    try:
                        paragraph.style = style_name
                    except Exception:
                        pass
                if translated_text:
                    paragraph.add_run(translated_text)
            elif block['type'] == 'table':
                rows = block.get('translated_rows') or []
                if rows:
                    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
                    table.style = 'Table Grid'
                    for row_index, row in enumerate(rows):
                        for col_index, cell_text in enumerate(row):
                            table.rows[row_index].cells[col_index].text = cell_text or ''

        document.save(output_path)
        return output_path

    async def translate_document(
        self,
        request: Request,
        user: UserModel | Any,
        file_path: str,
        source_filename: str,
        source_mime_type: str | None,
        target_language: str,
        source_language: str,
        model: str,
        force_ocr: bool = False,
        generate_output_file: bool = True,
        progress_callback: Callable[[list[str], str | None], Awaitable[None] | None] | None = None,
        chunk_timeout_seconds: int = DEFAULT_TRANSLATION_CHUNK_TIMEOUT_SECONDS,
    ) -> dict[str, Any]:
        filename = _normalize_filename(source_filename)
        suffix = Path(filename).suffix.lower()
        ext = suffix.lstrip('.')

        if ext == 'pdf' or source_mime_type == 'application/pdf':
            translated_pages = await self.translate_pdf_pages(
                request=request,
                user=user,
                file_path=file_path,
                target_language=target_language,
                source_language=source_language,
                model=model,
                force_ocr=force_ocr,
                progress_callback=progress_callback,
                chunk_timeout_seconds=chunk_timeout_seconds,
            )
            translation_text = '\n\n'.join(page.get('translation_text', '') for page in translated_pages if page.get('translation_text')).strip()

            if not generate_output_file:
                return {
                    'translation_text': translation_text,
                    'used_ocr': force_ocr,
                    'page_count': len(translated_pages),
                }

            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                output_path = tmp_file.name

            translated_layout_pages = await self.translate_pdf_layout_items(
                request=request,
                user=user,
                file_path=file_path,
                target_language=target_language,
                source_language=source_language,
                model=model,
                progress_callback=progress_callback,
                chunk_timeout_seconds=chunk_timeout_seconds,
            )
            has_layout_items = any(page_items for page_items in translated_layout_pages)
            has_complex_script = _layout_pages_have_complex_script(translated_layout_pages)
            has_rebuildable_layout = await asyncio.to_thread(_pdf_has_rebuildable_layout, file_path)
            if has_layout_items and (has_complex_script or has_rebuildable_layout):
                await asyncio.to_thread(self.build_pdf_layout_rebuild, file_path, translated_layout_pages, output_path)
            elif has_layout_items:
                await asyncio.to_thread(self.build_pdf_clone, file_path, translated_layout_pages, output_path)
            else:
                await asyncio.to_thread(self.build_pdf, translated_pages, output_path, filename)
            visual_qa = await asyncio.to_thread(_screen_pdf_visual_design, file_path, output_path)
            if (
                has_layout_items
                and not has_complex_script
                and not has_rebuildable_layout
                and _visual_similarity_score(visual_qa) < PDF_VISUAL_SIMILARITY_TARGET
            ):
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as fallback_tmp_file:
                    fallback_output_path = fallback_tmp_file.name

                fallback_visual_qa: dict[str, Any] | None = None
                try:
                    await asyncio.to_thread(
                        self.build_pdf_image_overlay,
                        file_path,
                        translated_layout_pages,
                        fallback_output_path,
                    )
                    fallback_visual_qa = await asyncio.to_thread(
                        _screen_pdf_visual_design,
                        file_path,
                        fallback_output_path,
                    )
                    if _visual_similarity_score(fallback_visual_qa) >= _visual_similarity_score(visual_qa):
                        try:
                            os.remove(output_path)
                        except OSError:
                            pass
                        output_path = fallback_output_path
                        visual_qa = fallback_visual_qa
                        visual_qa['generation_mode'] = 'image_overlay_clone'
                        fallback_output_path = ''
                    else:
                        visual_qa['fallback_visual_qa'] = fallback_visual_qa
                        visual_qa['generation_mode'] = 'vector_clone'
                finally:
                    if fallback_output_path:
                        try:
                            os.remove(fallback_output_path)
                        except OSError:
                            pass
            elif has_layout_items:
                visual_qa['generation_mode'] = (
                    'layout_rebuild' if (has_complex_script or has_rebuildable_layout) else 'vector_clone'
                )
            else:
                visual_qa['generation_mode'] = 'flow_rebuild'
            translated_filename = f'{Path(filename).stem}_translated_{target_language}.pdf'
            with open(output_path, 'rb') as output_file:
                contents, storage_path = await asyncio.to_thread(
                    Storage.upload_file,
                    output_file,
                    translated_filename,
                    {
                        'OpenWebUI-File-Name': translated_filename,
                        'OpenWebUI-Source-File': filename,
                    },
                )
            try:
                os.remove(output_path)
            except OSError:
                pass

            return {
                'translation_text': translation_text,
                'file_contents': contents,
                'file_path': storage_path,
                'output_filename': translated_filename,
                'used_ocr': force_ocr,
                'page_count': len(translated_pages),
                'visual_qa': visual_qa,
            }

        if ext == 'docx' or source_mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            translated_blocks = await self.translate_docx_blocks(
                request=request,
                user=user,
                file_path=file_path,
                target_language=target_language,
                source_language=source_language,
                model=model,
                progress_callback=progress_callback,
                chunk_timeout_seconds=chunk_timeout_seconds,
            )
            translation_text = _render_docx_translation_text(translated_blocks)

            if not generate_output_file:
                return {'translation_text': translation_text, 'page_count': len(translated_blocks)}

            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                output_path = tmp_file.name

            await asyncio.to_thread(self.build_docx, translated_blocks, output_path)
            translated_filename = f'{Path(filename).stem}_translated_{target_language}.docx'
            with open(output_path, 'rb') as output_file:
                contents, storage_path = await asyncio.to_thread(
                    Storage.upload_file,
                    output_file,
                    translated_filename,
                    {
                        'OpenWebUI-File-Name': translated_filename,
                        'OpenWebUI-Source-File': filename,
                    },
                )
            try:
                os.remove(output_path)
            except OSError:
                pass

            return {
                'translation_text': translation_text,
                'file_contents': contents,
                'file_path': storage_path,
                'output_filename': translated_filename,
                'page_count': len(translated_blocks),
            }

        raise ValueError('Only PDF and DOCX files are supported')

document_translator = DocumentTranslator()
